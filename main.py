import wikipedia
import re
from docx import Document
from translate import Translator

# reading the inputs from the user
name = input("Name: ")
language = input("Language: ")
title = input("Project title: ")

# mark that we want to use user chosen language
wikipedia.set_lang(language)

# search a page with user chosen title for project
while True:
    try:
        # if we found a page, we save it into wiki and break to stop the while
        wiki = wikipedia.page(title)
        break
    except:
        # if there is no page with user title, print message below and ask user to enter a new title
        title = input("Invalid project title, please choose another: ")

# we give to our text variable the content from wiki
text = wiki.content

# adjust our text
text = re.sub(r'==', '', text)
text = re.sub(r'\n', '\n    ', text)
# there we translate the end section indicator from wikipedia to user chosen language
translator= Translator(to_lang=language)
translation = translator.translate("See also")
# continue with the text adjust
split = text.split(translation, 1)
text = split[0]

# writing the results into the document and save it
document = Document()
paragraph = document.add_heading(title, 0)
paragraph.alignment = 1
document.add_paragraph('    ' + text)
paragraph = document.add_paragraph()
paragraph.alignment = 2
document.save(title + '.docx')